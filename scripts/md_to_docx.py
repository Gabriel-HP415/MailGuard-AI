"""Convert PhanCongThanhVien.md to .docx for printing/sharing."""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor


SRC = Path(__file__).resolve().parent.parent / "docs" / "PhanCongThanhVien.md"
DST = Path(__file__).resolve().parent.parent / "docs" / "PhanCongThanhVien.docx"


def parse_md(text: str) -> list[tuple[str, str]]:
    """Return list of (style, content). Styles:
    - 'h1' / 'h2' / 'h3' / 'h4'
    - 'p' (paragraph)
    - 'bullet' (list item)
    - 'code' (code block, will render as multiple 'code' rows)
    - 'table_row' (tuple of cells)
    - 'hr'
    """
    out: list[tuple[str, str]] = []
    lines = text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.rstrip()

        # Code block
        if stripped.startswith("```"):
            lang = stripped[3:].strip()
            block: list[str] = []
            i += 1
            while i < len(lines) and not lines[i].startswith("```"):
                block.append(lines[i])
                i += 1
            i += 1  # skip closing fence
            for code_line in block:
                out.append(("code", code_line))
            continue

        # Headings
        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m:
            level = len(m.group(1))
            out.append((f"h{level}", m.group(2)))
            i += 1
            continue

        # Horizontal rule
        if stripped == "---":
            out.append(("hr", ""))
            i += 1
            continue

        # Bullet
        m = re.match(r"^(\s*)[-*]\s+(.*)$", line)
        if m:
            indent = len(m.group(1)) // 2  # rough nesting
            out.append(("bullet", (indent, m.group(2))))
            i += 1
            continue

        # Table-like line (heuristic: | ... |)
        if stripped.startswith("|") and stripped.endswith("|"):
            row = [c.strip() for c in stripped.strip("|").split("|")]
            if re.match(r"^[-:\s|]+$", stripped):
                i += 1
                continue
            out.append(("table_row", row))
            i += 1
            continue

        # Blank line
        if not stripped:
            i += 1
            continue

        # Plain paragraph (collect until blank)
        paragraph_lines: list[str] = [stripped]
        i += 1
        while i < len(lines):
            nxt = lines[i].rstrip()
            if (
                not nxt
                or nxt.startswith("#")
                or nxt.startswith("```")
                or nxt.startswith("|")
                or re.match(r"^\s*[-*]\s+", nxt)
                or nxt == "---"
            ):
                break
            paragraph_lines.append(nxt)
            i += 1
        out.append(("p", " ".join(paragraph_lines)))
    return out


def add_runs(p, text: str) -> None:
    """Render **bold** and `code` inline markers."""
    token_re = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")
    parts = token_re.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.bold = True
            run.font.size = Pt(12)
        elif part.startswith("`") and part.endswith("`"):
            run = p.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(11)
        else:
            run = p.add_run(part)
            run.font.size = Pt(12)


def build_doc(toks: list[tuple[str, str]], doc: Document) -> None:
    table_buf: list[list[str]] = []

    def flush_table() -> None:
        if not table_buf:
            return
        rows = table_buf
        cols = max(len(r) for r in rows)
        for r in rows:
            while len(r) < cols:
                r.append("")
        tbl = doc.add_table(rows=len(rows), cols=cols)
        tbl.style = "Light Grid Accent 1"
        for r_idx, row in enumerate(rows):
            for c_idx, cell_val in enumerate(row):
                cell = tbl.rows[r_idx].cells[c_idx]
                cell.text = ""
                p = cell.paragraphs[0]
                add_runs(p, cell_val)
                if r_idx == 0:
                    for run in p.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    tc_pr = cell._tc.get_or_add_tcPr()
                    from docx.oxml import OxmlElement
                    from docx.oxml.ns import qn

                    shd = OxmlElement("w:shd")
                    shd.set(qn("w:fill"), "1D4ED8")
                    tc_pr.append(shd)
        table_buf.clear()

    for style, content in toks:
        if style == "h1":
            flush_table()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(10)
            run = p.add_run(content)
            run.bold = True
            run.font.size = Pt(20)
            run.font.color.rgb = RGBColor(0x1D, 0x4E, 0xD8)
        elif style == "h2":
            flush_table()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(content)
            run.bold = True
            run.font.size = Pt(15)
            run.font.color.rgb = RGBColor(0x1F, 0x2A, 0x44)
        elif style == "h3":
            flush_table()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(content)
            run.bold = True
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(0x37, 0x41, 0x51)
        elif style == "h4":
            flush_table()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(content)
            run.bold = True
            run.font.size = Pt(12)
        elif style == "p":
            flush_table()
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.3
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            add_runs(p, content)
        elif style == "bullet":
            indent, text_content = content
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(3)
            add_runs(p, text_content)
        elif style == "code":
            if not table_buf or table_buf[-1] != ["__code_want_table__"]:
                table_buf.append(["__code_want_table__"])
            tbl = doc.tables[-1] if table_buf == ["__code_want_table__"] else None
        elif style == "table_row":
            table_buf.append(list(content))
        elif style == "hr":
            flush_table()
            p = doc.add_paragraph()
            run = p.add_run("─" * 50)
            run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

    flush_table()


def render_code_block(doc: Document, code_lines: list[str]) -> None:
    """Render a contiguous block of code lines into a shaded table."""
    if not code_lines:
        return
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F1F5F9")
    tc_pr.append(shd)
    first = True
    for line in code_lines:
        if first:
            p = cell.paragraphs[0]
            first = False
        else:
            p = cell.add_paragraph()
        run = p.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(10)
    doc.add_paragraph()


def render_code(toks: list[tuple[str, str]], doc: Document) -> None:
    """Replace 'code' tokens into a single shaded table block."""
    new: list[tuple[str, str]] = []
    code_buf: list[str] = []
    for style, content in toks:
        if style == "code":
            code_buf.append(content)
        else:
            if code_buf:
                render_code_block(doc, code_buf)
                code_buf = []
            if style != "table_row":
                new.append((style, content))
    if code_buf:
        render_code_block(doc, code_buf)
    build_doc(new, doc)


def main() -> None:
    src = SRC.read_text(encoding="utf-8")
    toks = parse_md(src)
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.4)
        section.right_margin = Cm(2.4)
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), "Times New Roman")

    # Custom bullets that contain code/table need different handling — use
    # the simpler combined renderer which calls render_code_block for 'code'.
    render_code(toks, doc)
    doc.save(str(DST))
    import os
    print(f"[OK] Saved {DST}")
    print(f"     Size: {os.path.getsize(DST):,} bytes")


if __name__ == "__main__":
    main()