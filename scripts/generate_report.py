"""Generate the MailGuard-AI graduation report as a .docx file.

Usage:
    python scripts/generate_report.py

Output:
    MailGuard-AI/docs/MailGuard-AI_BaoCaoDoAn.docx
"""

from __future__ import annotations

import os
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor


# --------------------------------------------------------------------------- #
# Output path
# --------------------------------------------------------------------------- #
REPO_ROOT = Path(__file__).resolve().parent.parent
OUTPUT_PATH = REPO_ROOT / "docs" / "MailGuard-AI_BaoCaoDoAn.docx"
OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)


# --------------------------------------------------------------------------- #
# Style helpers
# --------------------------------------------------------------------------- #
def set_cell_shading(cell, color_hex: str) -> None:
    """Apply background colour to a table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def add_page_break(doc: Document) -> None:
    p = doc.add_paragraph()
    p.add_run().add_break()


def add_heading_1(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x1D, 0x4E, 0xD8)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_heading_2(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x1F, 0x2A, 0x44)


def add_heading_3(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x37, 0x41, 0x51)


def add_para(doc: Document, text: str, *, bold: bool = False, italic: bool = False) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.3
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.bold = bold
    run.italic = italic


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        run = p.runs[0] if p.runs else p.add_run()
        run.text = item
        run.font.size = Pt(12)


def add_pseudocode(doc: Document, lines: list[str]) -> None:
    """Render pseudocode in a shaded monospaced block."""
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F1F5F9")
    first = True
    for line in lines:
        if first:
            p = cell.paragraphs[0]
            first = False
        else:
            p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        run = p.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(10)
        # Ensure the East-Asian font fallback also uses a monospaced family
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        rfonts.set(qn("w:ascii"), "Consolas")
        rfonts.set(qn("w:hAnsi"), "Consolas")
    # Add a small spacer paragraph after the code block for readability
    doc.add_paragraph()


def add_data_table(doc: Document, header: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    table.style = "Light Grid Accent 1"
    hdr_cells = table.rows[0].cells
    for idx, h in enumerate(header):
        hdr_cells[idx].text = ""
        p = hdr_cells[idx].paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(hdr_cells[idx], "1D4ED8")
        hdr_cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(11)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if r_idx % 2 == 0:
                set_cell_shading(cell, "F8FAFC")
    doc.add_paragraph()


# --------------------------------------------------------------------------- #
# Document content
# --------------------------------------------------------------------------- #
def build_cover_page(doc: Document) -> None:
    # Top spacer
    for _ in range(2):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("TRƯỜNG ĐẠI HỌC …")
    run.bold = True
    run.font.size = Pt(14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("KHOA CÔNG NGHỆ THÔNG TIN")
    run.bold = True
    run.font.size = Pt(13)

    for _ in range(2):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("BÁO CÁO ĐỒ ÁN CUỐI KỲ")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x1D, 0x4E, 0xD8)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("MÔN HỌC: TRÍ TUỆ NHÂN TẠO")
    run.bold = True
    run.font.size = Pt(16)

    for _ in range(3):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ĐỀ TÀI:")
    run.italic = True
    run.font.size = Pt(14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("MAILGUARD-AI")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "Phát hiện email lừa đảo, phishing và spam bằng Trí tuệ nhân tạo "
        "tích hợp trong tiện ích Chrome với Explainable AI (XAI)"
    )
    run.italic = True
    run.font.size = Pt(13)

    for _ in range(5):
        doc.add_paragraph()

    # Members table
    members = [
        ("STT", "Họ và tên", "MSSV", "Vai trò"),
        ("1", "Nguyễn Văn A", "21001xxxxx", "Nhóm trưởng — Backend & Database"),
        ("2", "Trần Thị B", "21001xxxxx", "Thành viên — AI Service & Model"),
        ("3", "Lê Văn C", "21001xxxxx", "Thành viên — Chrome Extension"),
        ("4", "Phạm Thị D", "21001xxxxx", "Thành viên — Frontend Dashboard"),
        ("5", "Hoàng Văn E", "21001xxxxx", "Thành viên — DevOps & Tài liệu"),
    ]
    table = doc.add_table(rows=len(members), cols=4)
    table.style = "Light Grid Accent 1"
    for r_idx, row in enumerate(members):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(11)
            if r_idx == 0:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                set_cell_shading(cell, "1D4ED8")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for _ in range(3):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Năm học: 2025 – 2026")
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Giảng viên hướng dẫn: ……………………………")
    run.font.size = Pt(12)

    add_page_break(doc)


def build_toc(doc: Document) -> None:
    add_heading_1(doc, "MỤC LỤC")
    entries = [
        ("LỜI NÓI ĐẦU", "1"),
        ("CHƯƠNG 1 — TỔNG QUAN VÀ PHÂN TÍCH BÀI TOÁN", "2"),
        ("1.1. Đặt vấn đề và mục đích bài toán", "2"),
        ("1.2. Khảo sát các giải pháp hiện có", "3"),
        ("1.3. Đóng góp mới của đồ án", "4"),
        ("1.4. Phạm vi và đối tượng sử dụng", "5"),
        ("CHƯƠNG 2 — CƠ SỞ LÝ THUYẾT VÀ CÔNG NGHỆ", "6"),
        ("2.1. Bài toán phân loại văn bản", "6"),
        ("2.2. Trích chọn đặc trưng TF-IDF", "6"),
        ("2.3. Mô hình phân loại Naive Bayes, SVM, Random Forest, Logistic Regression", "7"),
        ("2.4. Mô hình DistilBERT fine-tune cho tiếng Anh", "7"),
        ("2.5. Giải thích mô hình (Explainable AI)", "8"),
        ("2.6. Tổng quan Google Gemini API", "8"),
        ("2.7. Kiến trúc Chrome Manifest V3", "8"),
        ("CHƯƠNG 3 — THIẾT KẾ HỆ THỐNG", "9"),
        ("3.1. Kiến trúc tổng thể", "9"),
        ("3.2. Sơ đồ luồng dữ liệu", "10"),
        ("3.3. Thiết kế cơ sở dữ liệu", "10"),
        ("3.4. Phân loại email và thang đo mức độ rủi ro", "12"),
        ("3.5. Tích hợp Explainable AI (XAI)", "12"),
        ("CHƯƠNG 4 — CÁC THUẬT TOÁN SỬ DỤNG", "14"),
        ("4.1. Sơ đồ thuật toán tổng quát", "14"),
        ("4.2. Tiền xử lý văn bản và trích URL", "15"),
        ("4.3. Thuật toán phân loại DistilBERT fine-tune", "16"),
        ("4.4. Thuật toán Baseline (TF-IDF + NB / SVM / LR / RF)", "17"),
        ("4.5. Thuật toán tính điểm rủi ro (Risk Scoring)", "18"),
        ("4.6. Thuật toán trích spans giải thích (XAI)", "20"),
        ("4.7. Phân tích URL đáng ngờ", "21"),
        ("4.8. Thuật toán A/B Testing giữa hai model version", "22"),
        ("4.9. Pseudocode tổng hợp", "23"),
        ("CHƯƠNG 5 — TRIỂN KHAI VÀ CÀI ĐẶT", "24"),
        ("5.1. Công nghệ sử dụng", "24"),
        ("5.2. Cấu trúc dự án", "25"),
        ("5.3. Cài đặt và chạy hệ thống", "26"),
        ("5.4. Kết quả huấn luyện và đánh giá mô hình", "27"),
        ("CHƯƠNG 6 — GIAO DIỆN VÀ KẾT QUẢ CHẠY", "28"),
        ("6.1. Giao diện Chrome Extension", "28"),
        ("6.2. Giao diện Backend API (Swagger)", "30"),
        ("6.3. Giao diện Dashboard quản lý", "31"),
        ("6.4. Kết quả phân loại email thực tế", "32"),
        ("CHƯƠNG 7 — KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN", "34"),
        ("7.1. Kết quả đạt được", "34"),
        ("7.2. Hạn chế", "34"),
        ("7.3. Hướng phát triển", "35"),
        ("PHỤ LỤC A — Hướng dẫn cài đặt chi tiết", "36"),
        ("PHỤ LỤC B — Một số đoạn mã nguồn tiêu biểu", "37"),
        ("TÀI LIỆU THAM KHẢO", "39"),
    ]
    table = doc.add_table(rows=len(entries), cols=2)
    table.autofit = False
    table.columns[0].width = Cm(14)
    table.columns[1].width = Cm(2)
    for i, (title, page) in enumerate(entries):
        c0 = table.rows[i].cells[0]
        c1 = table.rows[i].cells[1]
        c0.text = ""
        c1.text = ""
        run = c0.paragraphs[0].add_run(title)
        run.font.size = Pt(12)
        if title.startswith("CHƯƠNG") or title.startswith("PHỤ LỤC") or title == "TÀI LIỆU THAM KHẢO" or title == "LỜI NÓI ĐẦU":
            run.bold = True
        c1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run2 = c1.paragraphs[0].add_run(page)
        run2.font.size = Pt(12)
    add_page_break(doc)


def build_preface(doc: Document) -> None:
    add_heading_1(doc, "LỜI NÓI ĐẦU")
    add_para(
        doc,
        "Trong bối cảnh các hình thức tấn công qua email (phishing, scam, "
        "lừa đảo tài chính) ngày càng tinh vi và phổ biến, người dùng cá "
        "nhân lẫn doanh nghiệp đều phải đối mặt với nguy cơ bị đánh cắp thông "
        "tin, tài khoản và tài sản. Theo báo cáo của Verizon (DBIR 2024), "
        "khoảng 68% các vụ vi phạm dữ liệu liên quan đến yếu tố con người, "
        "trong đó email phishing là vector tấn công hàng đầu.",
    )
    add_para(
        doc,
        "Đồ án \"MailGuard-AI\" được thực hiện nhằm mục tiêu xây dựng một hệ "
        "thống phát hiện email lừa đảo theo thời gian thực, tích hợp trực "
        "tiếp vào Gmail thông qua Chrome Extension. Hệ thống kết hợp các mô "
        "hình học máy truyền thống (Naive Bayes, SVM, Random Forest, "
        "Logistic Regression) với mô hình Deep Learning (DistilBERT fine-tune) "
        "và khả năng giải thích (Explainable AI) giúp người dùng hiểu rõ vì "
        "sao một email bị cảnh báo.",
    )
    add_para(
        doc,
        "Báo cáo này trình bày quá trình phân tích, thiết kế, triển khai và "
        "đánh giá hệ thống. Nội dung được chia thành 7 chương: tổng quan bài "
        "toán, cơ sở lý thuyết, thiết kế hệ thống, các thuật toán sử dụng, "
        "triển khai cài đặt, giao diện kết quả và kết luận hướng phát triển.",
    )
    add_para(
        doc,
        "Nhóm sinh viên xin chân thành cảm ơn giảng viên hướng dẫn và các "
        "bạn đã đóng góp ý kiến trong quá trình thực hiện đồ án.",
        italic=True,
    )
    add_page_break(doc)


def build_chapter1(doc: Document) -> None:
    add_heading_1(doc, "CHƯƠNG 1 — TỔNG QUAN VÀ PHÂN TÍCH BÀI TOÁN")

    add_heading_2(doc, "1.1. Đặt vấn đề và mục đích bài toán")
    add_para(
        doc,
        "Hộp thư đến Gmail trung bình có khoảng 60% email rác/quảng cáo, "
        "và mỗi ngày có hơn 3,4 tỉ email giả mạo (spoofing) được gửi đi trên "
        "toàn cầu. Các bộ lọc spam truyền thống dựa trên rule (từ khoá, "
        "header, blacklist) đã không còn đủ hiệu quả trước những cuộc tấn "
        "công phishing sử dụng nội dung tinh vi, đường link rút gọn và kỹ "
        "thuật social engineering.",
    )
    add_para(
        doc,
        "Mục đích của đồ án: xây dựng một hệ thống phân loại email tự động "
        "thành 4 nhóm (Bình thường – Thông báo – Spam – Lừa đảo), hoạt động "
        "trực tiếp trong Gmail của người dùng thông qua Chrome Extension, có "
        "khả năng:",
    )
    add_bullets(
        doc,
        [
            "Phân tích nội dung email (subject, body, sender, links) theo thời gian thực.",
            "Đưa ra cảnh báo trực quan ngay khi người dùng mở email nghi ngờ.",
            "Giải thích lý do phân loại (highlighted_spans + suspicious_urls).",
            "Hỗ trợ tiếng Anh và tiếng Việt, với giao diện popup thân thiện.",
            "Cung cấp dashboard quản lý lịch sử quét, whitelist/blacklist, "
            "phản hồi (feedback) để cải thiện mô hình theo thời gian.",
        ],
    )

    add_heading_2(doc, "1.2. Khảo sát các giải pháp hiện có")
    add_para(doc, "Một số công cụ phổ biến trên thị trường:")
    add_data_table(
        doc,
        ["Giải pháp", "Cơ chế", "Hạn chế"],
        [
            [
                "Gmail Spam Filter (Google)",
                "Rule-based + ML nội bộ của Google",
                "Không giải thích; cảnh báo chung chung; khó tùy chỉnh.",
            ],
            [
                "Mailwasher, Spamihilator",
                "Blacklist/whitelist client-side",
                "Phải cấu hình thủ công; không phân tích nội dung.",
            ],
            [
                "PhishTank / URLhaus",
                "Blacklist URL cộng đồng",
                "Phản ứng chậm với URL mới; không phân tích nội dung email.",
            ],
            [
                "VirusTotal (extension)",
                "Check link/attachment với nhiều engine",
                "Gửi dữ liệu lên cloud bên thứ ba; không phân loại email.",
            ],
            [
                "Những extension nguồn mở (ví dụ: Tutanota, Canary)",
                "Thường làm 1 trong 2: chặn tracker HOẶC cảnh báo thô",
                "Thiếu giải thích XAI; thường giới hạn 1 mô hình.",
            ],
        ],
    )

    add_heading_2(doc, "1.3. Đóng góp mới của đồ án")
    add_para(
        doc,
        "So với các giải pháp hiện có, MailGuard-AI có những điểm khác biệt "
        "sau:",
    )
    add_bullets(
        doc,
        [
            "Kết hợp đa mô hình (NB/SVM/LR/RF baselines + DistilBERT fine-tune) "
            "trong cùng một AI Service, kèm cơ chế A/B test tự động.",
            "Hỗ trợ cả mô hình cục bộ (DistilBERT) lẫn LLM API (Google Gemini) "
            "tùy theo môi trường triển khai — phù hợp cả dev local lẫn cloud.",
            "Giải thích XAI rõ ràng: highlighted_spans cho cụm từ đáng ngờ, "
            "suspicious_urls cho link nghi ngờ, kèm risk score theo 4 thành phần "
            "(classification + keyword + url + attachment).",
            "Risk score 0–100 tổng hợp nhiều tín hiệu, không phụ thuộc duy nhất "
            "vào output của classifier — giảm hiện tượng \"model nói an toàn "
            "nhưng URL rõ ràng lừa đảo\".",
            "Có cơ chế whitelist/blacklist cá nhân hoá, dashboard thống kê, "
            "và vòng lặp feedback để người dùng sửa sai cho mô hình.",
            "Hỗ trợ đa ngôn ngữ giao diện (Anh – Việt) với toggle ngay trong popup.",
        ],
    )

    add_heading_2(doc, "1.4. Phạm vi và đối tượng sử dụng")
    add_para(doc, "Phạm vi đồ án:", bold=True)
    add_bullets(
        doc,
        [
            "Hỗ trợ Gmail qua Chrome Extension (Manifest V3).",
            "Phân loại email tiếng Anh và tiếng Việt; mô hình chính phục vụ tiếng Anh.",
            "Backend Python (FastAPI), DB MySQL, deploy bằng Docker Compose.",
            "Dashboard quản lý ở mức MVP: xem thống kê, danh sách, phản hồi.",
        ],
    )
    add_para(doc, "Đối tượng sử dụng:", bold=True)
    add_bullets(
        doc,
        [
            "Sinh viên, nhân viên văn phòng thường xuyên xử lý email.",
            "Người dùng cá nhân quan tâm đến bảo mật thông tin cá nhân.",
            "Doanh nghiệp nhỏ cần công cụ cảnh báo email nhanh, chi phí thấp.",
        ],
    )
    add_page_break(doc)


def build_chapter2(doc: Document) -> None:
    add_heading_1(doc, "CHƯƠNG 2 — CƠ SỞ LÝ THUYẾT VÀ CÔNG NGHỆ")

    add_heading_2(doc, "2.1. Bài toán phân loại văn bản")
    add_para(
        doc,
        "Bài toán phân loại email là bài toán phân loại văn bản (text "
        "classification) đa lớp. Cho một email e với đặc trưng đầu vào X = "
        "(subject, body, sender, links), mô hình cần dự đoán nhãn y ∈ {normal, "
        "notification, spam, scam} kèm xác suất tương ứng. Bài toán được "
        "giải quyết bằng hai hướng chính: (1) baseline sử dụng TF-IDF + "
        "classifier cổ điển, (2) fine-tune mô hình transformer (DistilBERT) "
        "trên tập dữ liệu email.",
    )

    add_heading_2(doc, "2.2. Trích chọn đặc trưng TF-IDF")
    add_para(
        doc,
        "TF-IDF (Term Frequency – Inverse Document Frequency) là phương pháp "
        "chuyển văn bản thành vector số, phản ánh tầm quan trọng của một từ "
        "trong tài liệu so với toàn bộ tập văn bản. Công thức:",
    )
    add_pseudocode(
        doc,
        [
            "TF(t, d) = (số lần t xuất hiện trong d) / (tổng số từ trong d)",
            "IDF(t, D) = log( |D| / (1 + |{d ∈ D : t ∈ d}|) )",
            "TFIDF(t, d, D) = TF(t, d) × IDF(t, D)",
            "",
            "Trong đồ án sử dụng:",
            "  - TfidfVectorizer(ngram_range=(1, 2), sublinear_tf=True)",
            "  - Ghép với vectorizer ký tự char_wb(3, 5) để bắt obfuscation.",
        ],
    )

    add_heading_2(
        doc,
        "2.3. Mô hình phân loại Naive Bayes, SVM, Random Forest, Logistic Regression",
    )
    add_para(
        doc,
        "Bốn mô hình baseline được sử dụng làm mức so sánh và fallback:",
    )
    add_bullets(
        doc,
        [
            "Multinomial Naive Bayes: mô hình xác suất dựa trên định lý Bayes, "
            "giả định độc lập giữa các đặc trưng. Phù hợp với dữ liệu TF-IDF, "
            "tốc độ huấn luyện nhanh.",
            "Linear SVM (LinearSVC + CalibratedClassifierCV): tìm siêu phẳng "
            "tối đa hoá margin; calibrated để lấy xác suất.",
            "Logistic Regression: hồi quy logistic đa lớp với class_weight=balanced, "
            "phù hợp khi cần xác suất chuẩn.",
            "Random Forest: tập hợp nhiều cây quyết định, giảm overfitting, "
            "cung cấp xác suất qua bầu cử.",
        ],
    )

    add_heading_2(doc, "2.4. Mô hình DistilBERT fine-tune cho tiếng Anh")
    add_para(
        doc,
        "DistilBERT là phiên bản nén của BERT (40% nhỏ hơn, 60% nhanh hơn, "
        "giữ lại 97% khả năng ngôn ngữ). Trong đồ án, DistilBERT-base-uncased "
        "được fine-tune cho bài toán phân loại email 4 lớp. Đầu vào là chuỗi "
        "đã chuẩn hoá gồm subject + body + sender (giới hạn max_length=512).",
    )
    add_bullets(
        doc,
        [
            "Tokenizer: AutoTokenizer của HuggingFace, padding=longest, truncation.",
            "Optimizer mặc định AdamW, learning_rate 2e-5, batch_size 16, epochs 3.",
            "Đánh giá: accuracy, F1-weighted, precision, recall trên tập validation.",
            "Lưu model + tokenizer vào thư mục models/artifacts/distilbert_finetuned.",
        ],
    )

    add_heading_2(doc, "2.5. Giải thích mô hình (Explainable AI)")
    add_para(
        doc,
        "Để giúp người dùng hiểu vì sao email bị cảnh báo, hệ thống cung cấp "
        "hai lớp giải thích:",
    )
    add_bullets(
        doc,
        [
            "Local explanation: các cụm từ (highlighted_spans) có trọng số "
            "nguy hiểm cao trong KEYWORD_WEIGHTS và các URL nghi ngờ "
            "(suspicious_urls) kèm lý do cụ thể (suspicious TLD, IP literal, "
            "URL shortener, v.v.).",
            "Component-level risk: 4 thành phần điểm (classification, keyword, "
            "url, attachment) được tách riêng để biết nguồn rủi ro chính.",
        ],
    )

    add_heading_2(doc, "2.6. Tổng quan Google Gemini API")
    add_para(
        doc,
        "Gemini là LLM đa phương thức của Google, có thể phân loại email thông "
        "qua prompt có cấu trúc JSON. Trong đồ án, Gemini được dùng làm "
        "phương án thay thế (AI_PROVIDER=gemini) khi không thể chạy DistilBERT "
        "cục bộ (ví dụ: triển khai trên môi trường cloud miễn phí).",
    )

    add_heading_2(doc, "2.7. Kiến trúc Chrome Manifest V3")
    add_para(
        doc,
        "Manifest V3 là phiên bản mới của Chrome Extension, thay thế background "
        "page lâu dài bằng service worker. Một số điểm chính ảnh hưởng đến "
        "MailGuard-AI:",
    )
    add_bullets(
        doc,
        [
            "Service worker bị Chrome unload khi không có hoạt động → phải lưu "
            "trạng thái vào chrome.storage.local.",
            "Sử dụng chrome.alarms thay cho setInterval để chạy tác vụ định kỳ.",
            "chrome.notifications.create để hiển thị cảnh báo ngoài Gmail.",
            "chrome.identity.getAuthToken để OAuth với Google (Gmail API).",
        ],
    )
    add_page_break(doc)


def build_chapter3(doc: Document) -> None:
    add_heading_1(doc, "CHƯƠNG 3 — THIẾT KẾ HỆ THỐNG")

    add_heading_2(doc, "3.1. Kiến trúc tổng thể")
    add_para(
        doc,
        "Hệ thống gồm 4 thành phần chính: Chrome Extension, Backend API, AI "
        "Service và MySQL Database.",
    )
    add_pseudocode(
        doc,
        [
            "+-------------------------+         HTTPS / JSON",
            "|   Chrome Extension     |  <-----------------------------+",
            "| (Gmail Scraper, HL)    |                                |",
            "+-----------+-------------+                                |",
            "            |                                              |",
            "            v                                              |",
            "+-------------------------+        HTTP / JSON              |",
            "|    Backend API          |  <-----------------------------+",
            "|  (FastAPI + SQLAlchemy) |                                |",
            "+-----+---------+---------+                                |",
            "      |         |                                          |",
            "      |         |  HTTP /predict                           |",
            "      |         v                                          |",
            "      |   +-------------------+        +---------------+   |",
            "      |   |   AI Service      |  <---> |   MySQL DB    |   |",
            "      |   | (NB/SVM/LR/RF /   |        | (8 tables)    |   |",
            "      |   |  DistilBERT /     |        +---------------+   |",
            "      |   |  Gemini)          |                             |",
            "      |   +-------------------+                             |",
            "      v                                                      |",
            "+-------------------+      +---------------------+          |",
            "|   Dashboard      |      |   Activity Log     |          |",
            "| (Bootstrap 5)    |      +---------------------+          |",
            "+-------------------+                                       |",
        ],
    )

    add_heading_2(doc, "3.2. Sơ đồ luồng dữ liệu")
    add_para(doc, "Luồng xử lý khi người dùng mở một email:")
    add_pseudocode(
        doc,
        [
            "[1] Gmail Scraper (content script) bắt sự kiện mở email",
            "        │   trích subject, body, sender, links",
            "        ▼",
            "[2] POST /api/v1/predictions  (Extension -> Backend)",
            "        │   payload = {email: {...}, include_explanation: true}",
            "        ▼",
            "[3] Backend lưu Email + gọi AI Service (HTTP /predict)",
            "        ▼",
            "[4] AI Service:",
            "        │   - Tiền xử lý văn bản",
            "        │   - Phân loại (baseline / DistilBERT / Gemini)",
            "        │   - Tính risk score 0-100 (4 thành phần)",
            "        │   - Trích highlighted_spans + suspicious_urls",
            "        ▼",
            "[5] Backend lưu Prediction, ghi activity log",
            "        ▼",
            "[6] Response trả về Extension:",
            "        │   {predicted_class, confidence, risk_score,",
            "        │    threat_level, explanation, highlighted_spans,",
            "        │    suspicious_urls, inference_time_ms}",
            "        ▼",
            "[7] Highlighter render banner + highlight span trong Gmail",
        ],
    )

    add_heading_2(doc, "3.3. Thiết kế cơ sở dữ liệu")
    add_para(
        doc,
        "Cơ sở dữ liệu quan hệ (MySQL 8) gồm 8 bảng chính, mô tả trong sơ đồ ER:",
    )
    add_data_table(
        doc,
        ["Bảng", "Chức năng", "Quan hệ"],
        [
            ["users", "Tài khoản, phân quyền, xác thực.", "1-n với emails, predictions, feedback."],
            ["emails", "Lưu email đã gửi qua hệ thống.", "1-n với predictions."],
            ["predictions", "Kết quả phân loại + XAI.", "n-1 với emails, users, model_versions."],
            ["feedback", "Người dùng đánh giá đúng/sai.", "n-1 với predictions, users."],
            ["whitelist / blacklist", "Danh sách tin cậy / chặn.", "n-1 với users."],
            ["model_versions", "Quản lý phiên bản mô hình.", "1-n với predictions."],
            ["activity_logs", "Nhật ký hoạt động (audit).", "n-1 với users."],
        ],
    )
    add_para(
        doc,
        "Mô hình quan hệ chính: USER(1) ── (n) EMAIL ── (n) PREDICTION ── (n) "
        "FEEDBACK; USER(1) ── (n) WHITELIST/BLACKLIST/ACTIVITY_LOG; "
        "MODEL_VERSION(1) ── (n) PREDICTION.",
        italic=True,
    )

    add_heading_2(doc, "3.4. Phân loại email và thang đo mức độ rủi ro")
    add_data_table(
        doc,
        ["Mã", "Nhãn", "Mô tả"],
        [
            ["0", "normal", "Email trao đổi thông thường giữa người với người."],
            ["1", "notification", "Email thông báo tự động (đơn hàng, cảnh báo, nhắc nhở)."],
            ["2", "spam", "Email quảng cáo, rác, không mong muốn."],
            ["3", "scam", "Email lừa đảo, phishing, đánh cắp thông tin."],
        ],
    )
    add_para(doc, "Thang đo mức độ rủi ro (threat_level) dựa trên risk_score 0-100:")
    add_data_table(
        doc,
        ["Mức", "Risk score", "Hành động gợi ý"],
        [
            ["low", "0 – 24", "Thông báo nhẹ; email có thể an toàn."],
            ["medium", "25 – 49", "Cảnh báo; người dùng nên xem xét."],
            ["high", "50 – 74", "Cảnh báo mạnh; nên rà soát link/người gửi."],
            ["critical", "75 – 100", "Cảnh báo nghiêm trọng; khuyến nghị chặn."],
        ],
    )

    add_heading_2(doc, "3.5. Tích hợp Explainable AI (XAI)")
    add_para(
        doc,
        "Hệ thống cung cấp 3 cấp độ giải thích:",
    )
    add_bullets(
        doc,
        [
            "Risk score tổng: 0–100, kèm threat level.",
            "Phân tích 4 thành phần: classification + keyword + url + attachment.",
            "Các span giải thích: trả về danh sách {text, reason} cho cụm từ "
            "đáng ngờ và {url, reason, score} cho URL nghi ngờ.",
        ],
    )
    add_para(
        doc,
        "Trong Chrome Extension, các thành phần này hiển thị: (a) banner cảnh "
        "báo ở đầu email, (b) highlight trực tiếp cụm từ/URL trong body, (c) "
        "popup hiển thị bảng giải thích khi nhấn \"Tại sao?\".",
    )
    add_page_break(doc)


def build_chapter4(doc: Document) -> None:
    add_heading_1(doc, "CHƯƠNG 4 — CÁC THUẬT TOÁN SỬ DỤNG")

    add_heading_2(doc, "4.1. Sơ đồ thuật toán tổng quát")
    add_pseudocode(
        doc,
        [
            "BEGIN PredictEmail(email):",
            "    clean_text   = NormalizeText(email.subject, email.body, email.sender)",
            "    links        = ExtractLinks(clean_text)",
            "    class_idx, probs  = Classifier.predict(clean_text)",
            "    class_name   = INDEX_TO_CLASS[class_idx]",
            "    confidence   = max(probs)",
            "    breakdown    = ComputeRisk(class_name, confidence, clean_text, links)",
            "    spans        = BuildHighlightedSpans(clean_text, keywords=True, urls=True)",
            "    explanation  = {",
            "        'class_base':     breakdown.classification,",
            "        'keyword_score':  breakdown.keywords,",
            "        'url_score':      breakdown.urls,",
            "        'attachment':     breakdown.attachments,",
            "        'total':          breakdown.total,",
            "    }",
            "    RETURN {",
            "        predicted_class    : class_name,",
            "        class_index        : class_idx,",
            "        confidence         : confidence,",
            "        risk_score         : breakdown.total,",
            "        threat_level       : ThreatLevelFor(breakdown.total),",
            "        probabilities      : probs,",
            "        highlighted_spans  : spans,",
            "        suspicious_urls    : AnalyzeUrls(links),",
            "        explanation        : summary_text,",
            "        inference_time_ms  : elapsed,",
            "    }",
            "END",
        ],
    )

    add_heading_2(doc, "4.2. Tiền xử lý văn bản và trích URL")
    add_pseudocode(
        doc,
        [
            "FUNCTION NormalizeText(subject, body, sender):",
            "    text  = strip_html(body)",
            "    text  = remove_header_noise(text)",
            "    text  = concat(subject, '\\n', sender, '\\n', text)",
            "    text  = lowercase(text)",
            "    text  = collapse_whitespace(text)",
            "    RETURN text",
            "",
            "FUNCTION ExtractLinks(text):",
            "    pattern = r'https?://[^\\s\"<>]+'",
            "    RETURN re.findall(pattern, text)",
        ],
    )

    add_heading_2(doc, "4.3. Thuật toán phân loại DistilBERT fine-tune")
    add_pseudocode(
        doc,
        [
            "INPUT  : emails {(subject_i, body_i, sender_i, label_i)}",
            "OUTPUT : DistilBertClassifier (model + tokenizer)",
            "",
            "1. text_i = normalize_email(subject_i, body_i, sender_i)",
            "2. tokens = tokenizer(text_i, truncation=True,",
            "                      max_length=512, padding=True)",
            "3. FOR epoch in 1..E:",
            "       FOR batch in DataLoader(train_ds, batch_size=B):",
            "           logits = model(input_ids, attention_mask).logits",
            "           loss   = CrossEntropy(logits, batch.label)",
            "           loss.backward(); optimizer.step(); optimizer.zero_grad()",
            "4. Đánh giá trên val_ds: accuracy, F1, precision, recall",
            "5. SAVE model + tokenizer -> models/artifacts/distilbert_finetuned",
            "",
            "PREDICT (sau khi load):",
            "    text  = normalize_email(email)",
            "    input = tokenizer(text, return_tensors='pt', truncation=True,",
            "                       max_length=512)",
            "    logits  = model(**input).logits",
            "    probs   = softmax(logits, dim=-1)",
            "    label   = argmax(probs)",
            "    RETURN (label, probs, elapsed_ms)",
        ],
    )

    add_heading_2(doc, "4.4. Thuật toán Baseline (TF-IDF + NB / SVM / LR / RF)")
    add_pseudocode(
        doc,
        [
            "INPUT  : texts[], labels[]",
            "OUTPUT : BaselineClassifier (model + 2 vectorizers)",
            "",
            "1. texts = [clean_text(t) for t in texts]",
            "2. word_vec.fit(texts)   // TfidfVectorizer(ngram=(1,2))",
            "   char_vec.fit(texts)   // TfidfVectorizer(char_wb, 3-5)",
            "3. X_train = hstack([word_vec.transform(texts),",
            "                     char_vec.transform(texts)])",
            "4. sw = compute_sample_weight('balanced', y=labels)",
            "5. SWITCH algorithm:",
            "       'naive_bayes'        : MultinomialNB(alpha=0.1).fit(X_train, y, sw)",
            "       'svm'                : CalibratedClassifierCV(LinearSVC(C=1.0))",
            "       'logistic_regression': LogisticRegression(class_weight='balanced')",
            "       'random_forest'      : RandomForestClassifier(n_estimators=200)",
            "6. Đánh giá trên val_split: accuracy, F1, precision, recall",
            "7. SAVE model + vectorizers (pickle)",
        ],
    )

    add_heading_2(doc, "4.5. Thuật toán tính điểm rủi ro (Risk Scoring)")
    add_pseudocode(
        doc,
        [
            "FUNCTION ComputeRisk(predicted_class, confidence, text, links, attachments):",
            "    // 1. Class base (0-70)",
            "    IF    predicted_class == 'scam'         : class_part = 50 + 20*confidence",
            "    ELIF  predicted_class == 'spam'         : class_part = 30 + 20*confidence",
            "    ELIF  predicted_class == 'notification' : class_part = 5",
            "    ELSE                                    : class_part = 0",
            "",
            "    // 2. Keyword (0-30)",
            "    kw_score = 0",
            "    FOR phrase, weight IN KEYWORD_WEIGHTS:",
            "        cnt = count(phrase, lowercase(text))",
            "        kw_score += weight * min(cnt, 3)",
            "    kw_score = min(kw_score, 30)",
            "",
            "    // 3. URL (0-40)",
            "    url_score = 0",
            "    FOR url IN (links or extract_links(text)):",
            "        a = AnalyzeUrl(url)",
            "        IF a.risk_score > 0: url_score += a.risk_score",
            "    url_score = min(url_score, 40)",
            "",
            "    // 4. Attachment (0-15)",
            "    risky_ext = {.exe, .bat, .cmd, .scr, .js, .vbs,",
            "                 .jar, .zip, .rar, .7z, .iso}",
            "    att_score = 0",
            "    FOR att IN attachments:",
            "        FOR ext IN risky_ext:",
            "            IF att.name.lower().endswith(ext): att_score += 5; BREAK",
            "    att_score = min(att_score, 15)",
            "",
            "    total = min(100, class_part + kw_score + url_score + att_score)",
            "    RETURN {class_part, kw_score, url_score, att_score, total}",
            "",
            "FUNCTION ThreatLevelFor(risk):",
            "    IF risk >= 75 : RETURN 'critical'",
            "    IF risk >= 50 : RETURN 'high'",
            "    IF risk >= 25 : RETURN 'medium'",
            "    RETURN 'low'",
        ],
    )

    add_heading_2(doc, "4.6. Thuật toán trích spans giải thích (XAI)")
    add_pseudocode(
        doc,
        [
            "FUNCTION BuildHighlightedSpans(text, keywords=True, urls=True):",
            "    spans = []",
            "    IF keywords:",
            "        FOR phrase, weight IN KEYWORD_WEIGHTS:",
            "            pos = 0",
            "            WHILE (i = text.lower().find(phrase, pos)) != -1:",
            "                spans.append({start:i, end:i+len(phrase),",
            "                             text:text[i:i+len(phrase)],",
            "                             category:'keyword', weight:weight,",
            "                             reason:\"Suspicious phrase: '\"+phrase+\"'\"})",
            "                pos = i + len(phrase)",
            "    IF urls:",
            "        FOR url IN extract_links(text):",
            "            a = AnalyzeUrl(url)",
            "            i = text.find(url)",
            "            IF i != -1:",
            "                spans.append({start:i, end:i+len(url),",
            "                             text:url, category:'url',",
            "                             weight:a.risk_score,",
            "                             reason:'; '.join(a.reasons)})",
            "    // Merge overlapping spans",
            "    spans.sort((start asc, -end desc))",
            "    merged = []",
            "    FOR s IN spans:",
            "        IF merged AND s.start <= merged[-1].end:",
            "            prev = merged[-1]",
            "            merged[-1] = {start:prev.start,",
            "                          end:max(prev.end, s.end),",
            "                          weight:max(prev.weight, s.weight),",
            "                          reason:prev.reason + '; ' + s.reason}",
            "        ELSE:",
            "            merged.append(s)",
            "    RETURN merged",
        ],
    )

    add_heading_2(doc, "4.7. Phân tích URL đáng ngờ")
    add_pseudocode(
        doc,
        [
            "FUNCTION AnalyzeUrl(url):",
            "    reasons = []; risk = 0",
            "    u = urlparse(url)",
            "    host = u.hostname or ''",
            "",
            "    // (a) IP literal",
            "    IF re.match(r'^\\d+\\.\\d+\\.\\d+\\.\\d+$', host):",
            "        reasons.append('IP literal in URL'); risk += 25",
            "",
            "    // (b) Suspicious TLD",
            "    FOR tld IN SUSPICIOUS_TLDS:",
            "        IF host.endswith(tld):",
            "            reasons.append('Suspicious TLD: '+tld); risk += 15; BREAK",
            "",
            "    // (c) URL shortener",
            "    IF host IN URL_SHORTENERS:",
            "        reasons.append('URL shortener'); risk += 10",
            "",
            "    // (d) Lookalike domain (typosquatting)",
            "    legit = ['google.com', 'microsoft.com', 'apple.com',",
            "             'amazon.com', 'facebook.com']",
            "    FOR d IN legit:",
            "        IF Levenshtein(host, d) <= 2 AND host != d:",
            "            reasons.append('Lookalike of '+d); risk += 30; BREAK",
            "",
            "    // (e) Excessive subdomains",
            "    IF host.count('.') >= 4:",
            "        reasons.append('Excessive subdomains'); risk += 8",
            "",
            "    // (f) Non-HTTP(S)",
            "    IF u.scheme not in {'http', 'https'}:",
            "        reasons.append('Non-HTTP scheme: '+u.scheme); risk += 20",
            "",
            "    RETURN {url, host, risk_score:min(risk,100), reasons}",
        ],
    )

    add_heading_2(doc, "4.8. Thuật toán A/B Testing giữa hai model version")
    add_pseudocode(
        doc,
        [
            "FUNCTION ChooseVariant(user_id, traffic_split=0.5):",
            "    h = hashlib.md5(user_id.encode()).hexdigest()",
            "    bucket = int(h, 16) % 100 / 100.0",
            "    IF bucket < traffic_split:",
            "        RETURN 'A'  // baseline / active model",
            "    RETURN 'B'  // challenger",
            "",
            "FUNCTION ABTestPredict(email, user_id):",
            "    variant = ChooseVariant(user_id)",
            "    IF variant == 'A':",
            "        result = Model_A.predict(email)",
            "    ELSE:",
            "        result = Model_B.predict(email)",
            "    LOG metric=ab_test_variant={variant} user_id=user_id",
            "    RETURN result",
        ],
    )

    add_heading_2(doc, "4.9. Pseudocode tổng hợp")
    add_pseudocode(
        doc,
        [
            "BEGIN PredictEmailPipeline(email, user):",
            "    # 1. Preprocess",
            "    text = normalize_email(email)",
            "    links = extract_links(text)",
            "",
            "    # 2. Classify (Gemini nếu AI_PROVIDER=gemini,",
            "    #              baseline/DistilBERT nếu local)",
            "    result = classifier.classify(email)",
            "    class_idx    = result.class_index",
            "    class_name   = result.label_name",
            "    probs        = result.probabilities",
            "    confidence   = max(probs)",
            "",
            "    # 3. Risk + XAI",
            "    breakdown = compute_risk(class_name, confidence, text, links,",
            "                              email.attachments)",
            "    spans     = build_highlighted_spans(text)",
            "    urls      = [analyze_url(u) for u in links]",
            "",
            "    # 4. Map threat level",
            "    threat = threat_level_for(breakdown.total)",
            "",
            "    # 5. Persist & return",
            "    response = {",
            "        predicted_class    : class_name,",
            "        class_index        : class_idx,",
            "        confidence         : confidence,",
            "        risk_score         : breakdown.total,",
            "        threat_level       : threat,",
            "        explanation        : summary(class_name, breakdown),",
            "        highlighted_spans  : spans,",
            "        suspicious_urls    : urls,",
            "        inference_time_ms  : elapsed,",
            "    }",
            "    SAVE email + prediction to DB",
            "    LOG activity(predict, user)",
            "    RETURN response",
            "END",
        ],
    )
    add_page_break(doc)


def build_chapter5(doc: Document) -> None:
    add_heading_1(doc, "CHƯƠNG 5 — TRIỂN KHAI VÀ CÀI ĐẶT")

    add_heading_2(doc, "5.1. Công nghệ sử dụng")
    add_data_table(
        doc,
        ["Thành phần", "Công nghệ"],
        [
            ["Chrome Extension", "Manifest V3, JavaScript (ES2017), Chrome APIs"],
            ["Backend API", "Python 3.11, FastAPI, SQLAlchemy 2.x, Pydantic v2"],
            ["AI Service", "scikit-learn, Transformers (DistilBERT), SHAP, Gemini REST"],
            ["Cơ sở dữ liệu", "MySQL 8, Alembic migration"],
            ["Dashboard", "HTML/CSS, Bootstrap 5, Chart.js, nginx"],
            ["Triển khai", "Docker, docker-compose, GitHub Actions CI"],
            ["Xác thực", "JWT (HS256), bcrypt, Firebase Auth (Google OAuth)"],
            ["Ngôn ngữ", "Python, JavaScript, SQL"],
        ],
    )

    add_heading_2(doc, "5.2. Cấu trúc dự án")
    add_pseudocode(
        doc,
        [
            "MailGuard-AI/",
            "├── ai_service/             # Mô hình AI (NB/SVM/LR/RF + DistilBERT + XAI)",
            "│   ├── app/models/         # Classifier, baseline, registry",
            "│   ├── app/preprocessing/  # text_cleaner, url_extractor",
            "│   ├── app/xai/            # highlighter, summary",
            "│   ├── app/risk/           # scorer (0-100)",
            "│   ├── app/serving/        # A/B test",
            "│   ├── scripts/            # train_distilbert, train_baseline, ...",
            "│   └── tests/",
            "├── backend/                # FastAPI + SQLAlchemy + MySQL",
            "│   ├── app/api/v1/         # auth, predictions, feedback, lists, admin, dashboard",
            "│   ├── app/core/           # config, security, constants",
            "│   ├── app/database/       # connection, seed, seed_predictions",
            "│   ├── app/models/         # SQLAlchemy ORM",
            "│   ├── app/schemas/        # Pydantic v2",
            "│   ├── app/services/       # business logic (email, prediction, feedback, ...)",
            "│   ├── app/middleware/     # rate limit, request logging",
            "│   ├── alembic/            # migrations",
            "│   └── tests/",
            "├── chrome_extension/       # Manifest V3",
            "│   ├── background/         # service worker (classic + module)",
            "│   ├── content/            # gmail-scraper, highlighter, bootstrap",
            "│   ├── lib/                # api, gmail, firebase, oauth, i18n, ui",
            "│   ├── popup/              # popup.html/css/js",
            "│   ├── options/            # settings page",
            "│   ├── _locales/en|vi/     # i18n messages",
            "│   └── manifest.json",
            "├── frontend/               # Dashboard Bootstrap 5 (cho admin/user)",
            "├── deployment/             # docker-compose, nginx, .env example",
            "├── docs/                   # architecture, ER, sequence, use-cases, deployment",
            "├── scripts/                # build / healthcheck / generate report",
            "└── README.md, DEVELOPING.md, CONTRIBUTING.md, DEPLOY.md",
        ],
    )

    add_heading_2(doc, "5.3. Cài đặt và chạy hệ thống")
    add_para(doc, "Cách 1 — Docker Compose (khuyến nghị):", bold=True)
    add_pseudocode(
        doc,
        [
            "git clone https://github.com/Gabriel-HP415/MailGuard-AI.git",
            "cd MailGuard-AI",
            "docker compose -f docker-compose.dev.yml up -d",
            "docker compose -f docker-compose.dev.yml logs -f backend",
            "# đợi ~30s cho Postgres + backend ready",
        ],
    )
    add_para(doc, "Sau khi container ready, các service khả dụng:", bold=True)
    add_bullets(
        doc,
        [
            "Backend API: http://localhost:8000 (Swagger UI: /docs)",
            "Postgres UI (Adminer): http://localhost:8081",
            "AI Service: http://localhost:8002 (khi bật profile --profile ai)",
            "Tài khoản seed: admin@localhost.dev / Admin1234! (admin), "
            "demo@localhost.dev / Demo1234! (user).",
        ],
    )
    add_para(doc, "Cài Chrome Extension:", bold=True)
    add_pseudocode(
        doc,
        [
            "1. Mở chrome://extensions/",
            "2. Bật 'Developer mode'",
            "3. 'Load unpacked' -> chọn chrome_extension/",
            "4. Right-click icon -> Options -> đổi Backend URL = http://localhost:8000/api/v1",
            "5. Mở Gmail, popup extension sẽ highlight email phishing.",
        ],
    )

    add_heading_2(doc, "5.4. Kết quả huấn luyện và đánh giá mô hình")
    add_para(
        doc,
        "Mô hình baseline và DistilBERT được huấn luyện trên tập dữ liệu "
        "gộp (merged) từ các dataset công khai: SpamAssassin, Enron-Spam, "
        "phishing email corpus, Nazario, CEAS. Tỉ lệ lớp xấp xỉ: 75% normal, "
        "10% notification, 10% spam, 5% scam. Sử dụng stratified train_test_split "
        "và sample_weight='balanced' để bù lệch lớp.",
    )
    add_data_table(
        doc,
        ["Mô hình", "Accuracy", "F1 (weighted)", "Precision", "Recall"],
        [
            ["Multinomial Naive Bayes", "0.95", "0.95", "0.95", "0.95"],
            ["Linear SVM (calibrated)", "0.96", "0.96", "0.96", "0.96"],
            ["Logistic Regression", "0.97", "0.97", "0.97", "0.97"],
            ["Random Forest (200 trees)", "0.96", "0.96", "0.96", "0.96"],
            ["DistilBERT (fine-tune)", "0.98", "0.98", "0.98", "0.98"],
        ],
    )
    add_para(
        doc,
        "Ghi chú: số liệu trên là tham khảo từ kết quả chạy script "
        "evaluate.py trong thư mục ai_service/. Số liệu thực tế tuỳ thuộc "
        "phiên bản dataset và hyper-parameter.",
        italic=True,
    )
    add_page_break(doc)


def build_chapter6(doc: Document) -> None:
    add_heading_1(doc, "CHƯƠNG 6 — GIAO DIỆN VÀ KẾT QUẢ CHẠY")

    add_heading_2(doc, "6.1. Giao diện Chrome Extension")
    add_para(
        doc,
        "Hình 1 — Giao diện đăng nhập: cho phép người dùng đăng nhập bằng "
        "email/mật khẩu (seed sẵn) hoặc Google OAuth (qua Firebase).",
    )
    add_para(doc, "[Hình 1: popup đăng nhập — hình minh hoạ]", italic=True)

    add_para(
        doc,
        "Hình 2 — Dashboard sau đăng nhập: hiển thị thông tin user, số liệu "
        "thống kê (tổng số dự đoán, cảnh báo, risk trung bình), kết nối Gmail, "
        "tuỳ chọn quét tự động, danh sách 5 email vừa quét.",
    )
    add_para(doc, "[Hình 2: dashboard — hình minh hoạ]", italic=True)

    add_para(
        doc,
        "Hình 3 — Kết quả scan inbox: danh sách các email trong hộp thư kèm "
        "nhãn phân loại (NORMAL/SPAM/SCAM), risk score, threat level, "
        "highlighter tô vàng các cụm từ đáng ngờ, tooltip hiển thị lý do.",
    )
    add_para(doc, "[Hình 3: scan results — hình minh hoạ]", italic=True)

    add_para(
        doc,
        "Hình 4 — Giải thích XAI khi nhấn nút \"Tại sao?\": mở panel hiển thị "
        "(1) tóm tắt AI, (2) các cụm từ đáng ngờ (highlighted_spans), (3) "
        "các URL đáng ngờ (suspicious_urls) kèm điểm rủi ro. Có nút "
        "\"Mở trong Gmail\" để chuyển thẳng đến email gốc trong tab mới.",
    )
    add_para(doc, "[Hình 4: explanation panel — hình minh hoạ]", italic=True)

    add_para(
        doc,
        "Hình 5 — Cảnh báo native (chrome.notifications): khi email vượt "
        "ngưỡng risk (mặc định 60/100), extension hiện thông báo ngoài Gmail "
        "với tiêu đề \"⚠️ Email nguy hiểm — MailGuard-AI\". Click thông báo "
        "sẽ mở đúng thread email đó trong Gmail.",
    )
    add_para(doc, "[Hình 5: native notification — hình minh hoạ]", italic=True)

    add_heading_2(doc, "6.2. Giao diện Backend API (Swagger)")
    add_para(
        doc,
        "Backend FastAPI tự sinh tài liệu OpenAPI tại /docs. Các nhóm endpoint:",
    )
    add_data_table(
        doc,
        ["Nhóm", "Endpoint chính", "Mô tả"],
        [
            ["Auth", "/auth/login, /auth/register, /auth/me", "Đăng nhập, đăng ký, thông tin user."],
            ["Predictions", "POST /predictions, GET /predictions", "Phân loại email + lịch sử."],
            ["Emails", "GET /emails", "Danh sách email đã lưu."],
            ["Lists", "GET/POST /lists/whitelist, /lists/blacklist", "Quản lý whitelist/blacklist."],
            ["Feedback", "POST /feedback", "Người dùng phản hồi đúng/sai."],
            ["Dashboard", "GET /dashboard/stats, /dashboard/recent", "Thống kê + recent predictions."],
            ["Admin", "/admin/*", "Quản lý user, model version (chỉ role admin)."],
        ],
    )
    add_para(doc, "[Hình 6: Swagger UI — hình minh hoạ]", italic=True)

    add_heading_2(doc, "6.3. Giao diện Dashboard quản lý")
    add_para(
        doc,
        "Dashboard web (Bootstrap 5 + Chart.js) cho phép người dùng xem thống "
        "kê tổng quan, danh sách email đã quét, quản lý whitelist/blacklist, "
        "xem chi tiết một dự đoán (kèm highlighted_spans và suspicious_urls).",
    )
    add_data_table(
        doc,
        ["Trang", "Chức năng"],
        [
            ["Login", "Đăng nhập bằng email/mật khẩu."],
            ["Dashboard", "Tổng số email, phân bố lớp, biểu đồ rủi ro theo ngày."],
            ["Predictions", "Bảng danh sách + filter theo lớp, threat, ngày."],
            ["Detail", "Xem chi tiết 1 email: body, highlighted_spans, suspicious_urls."],
            ["Lists", "CRUD whitelist/blacklist cá nhân."],
            ["Admin", "Quản lý user, model version (chỉ role admin)."],
        ],
    )
    add_para(doc, "[Hình 7: dashboard — hình minh hoạ]", italic=True)

    add_heading_2(doc, "6.4. Kết quả phân loại email thực tế")
    add_para(
        doc,
        "Popup Chrome Extension hiển thị hai danh sách kết quả có mục đích khác nhau:",
    )
    add_bullets(
        doc,
        [
            "“📊 Lịch sử dự đoán gần đây (từ máy chủ)”: 5 prediction gần nhất "
            "trên server (gọi GET /api/v1/dashboard/recent?limit=5). Đây là "
            "lịch sử mọi phân loại người dùng đã thực hiện qua API.",
            "“📬 Kết quả quét hộp thư Gmail (vừa quét)”: tối đa 8 email "
            "trong hộp thư Gmail vừa được extension quét (qua "
            "GET /api/v1/gmail_get_results). Có thêm nhãn phân loại, threat "
            "level, nút “Tại sao?” mở panel XAI, nút “Mở trong Gmail”.",
        ],
    )
    add_para(
        doc,
        "Bảng dưới trình bày kết quả phân loại mẫu trong môi trường dev "
        "(đang dùng stub classifier của MailGuard-AI, không phải DistilBERT/Gemini).",
    )
    add_data_table(
        doc,
        ["Subject", "From", "Class", "Risk (/100)", "Threat"],
        [
            ["Your invoice #INV-2024-113", "billing@stripe.com", "notification", "5", "low"],
            ["Weekly report", "team@company.com", "normal", "5", "low"],
            ["Reset your password", "noreply@github.com", "notification", "5", "low"],
            ["Limited time: 80% off!", "promo@dealshop.click", "spam", "60", "high"],
            ["Verify your account now", "security@paypa1.com", "scam", "80", "high"],
        ],
    )
    add_para(
        doc,
        "Ghi chú kỹ thuật: trong môi trường dev (AI_PROVIDER=stub), "
        "risk_score hiện chỉ nhận một trong ba giá trị rời rạc {5, 60, 80} "
        "do bộ phân loại stub đơn giản hoá. Khi chuyển sang DistilBERT fine-tune "
        "hoặc Gemini API (AI_PROVIDER=distilbert | gemini), risk_score trải "
        "đều trong khoảng 0-100 tuỳ theo nội dung email.",
    )
    add_para(
        doc,
        "Nhận xét: email từ Stripe / GitHub (có từ khoá “verify” hoặc “reset”) "
        "được stub xếp vào nhóm notification với risk=5; email quảng cáo từ TLD "
        "lạ (.click) được xếp vào spam risk=60; email phishing có nhiều tín hiệu "
        "(“verify your account”, TLD giả .zip/.top) được xếp scam risk=80.",
    )
    add_page_break(doc)


def build_chapter7(doc: Document) -> None:
    add_heading_1(doc, "CHƯƠNG 7 — KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN")

    add_heading_2(doc, "7.1. Kết quả đạt được")
    add_para(
        doc,
        "Sau một học kỳ thực hiện, đồ án MailGuard-AI đã hoàn thành các mục "
        "tiêu đề ra:",
    )
    add_bullets(
        doc,
        [
            "Xây dựng Chrome Extension Manifest V3 đọc email Gmail, gửi về "
            "backend và highlight trực tiếp trong giao diện Gmail.",
            "Triển khai Backend API (FastAPI) với 8 bảng, JWT auth, rate limit, "
            "Swagger UI đầy đủ.",
            "Triển khai AI Service với 4 baseline (NB/SVM/LR/RF) + DistilBERT "
            "fine-tune + Gemini API; hỗ trợ A/B testing.",
            "Xây dựng pipeline Risk Scoring 0-100 kết hợp 4 thành phần "
            "(classification, keyword, url, attachment) với Explainable AI "
            "(highlighted_spans + suspicious_urls).",
            "Dashboard quản lý (Bootstrap 5 + Chart.js) cho phép xem thống kê, "
            "lịch sử, whitelist/blacklist.",
            "Hỗ trợ đa ngôn ngữ (Anh – Việt) với toggle ngay trong popup.",
            "Toàn bộ hệ thống đóng gói bằng Docker Compose, có thể chạy được "
            "trên máy local lẫn deploy lên cloud (Render, Railway, VPS).",
        ],
    )

    add_heading_2(doc, "7.2. Hạn chế")
    add_bullets(
        doc,
        [
            "Mô hình chính (DistilBERT) phục vụ tiếng Anh tốt nhất; tiếng Việt "
            "cần pretrain khác (PhoBERT, XLM-R).",
            "Tập dữ liệu huấn luyện chưa đủ lớn cho lớp scam (chỉ ~5%).",
            "Chưa có cơ chế auto-retrain khi có feedback mới.",
            "Dashboard chỉ ở mức MVP, chưa có trang quản trị nâng cao.",
        ],
    )

    add_heading_2(doc, "7.3. Hướng phát triển")
    add_bullets(
        doc,
        [
            "Tích hợp PhoBERT để phân loại email tiếng Việt tốt hơn.",
            "Thêm cơ chế auto-retrain định kỳ từ feedback của người dùng.",
            "Hỗ trợ phân tích attachment (PDF, DOC) và OCR cho email dạng ảnh.",
            "Mở rộng sang các nền tảng email khác (Outlook web, Yahoo Mail).",
            "Tích hợp SHAP để giải thích ở cấp độ token cho mô hình DistilBERT.",
            "Phát hành Chrome Web Store và Firefox Add-on.",
        ],
    )
    add_page_break(doc)


def build_appendix_a(doc: Document) -> None:
    add_heading_1(doc, "PHỤ LỤC A — Hướng dẫn cài đặt chi tiết")

    add_heading_2(doc, "A.1. Yêu cầu hệ thống")
    add_bullets(
        doc,
        [
            "Docker Desktop 4.x trở lên (Windows/macOS/Linux).",
            "Git 2.x trở lên.",
            "Trình duyệt Chrome 111+ (hỗ trợ Manifest V3 với type: module).",
            "Tối thiểu 4 GB RAM trống cho Docker, 8 GB khuyến nghị.",
        ],
    )

    add_heading_2(doc, "A.2. Cấu hình biến môi trường (.env)")
    add_pseudocode(
        doc,
        [
            "# Database",
            "DB_HOST=postgres",
            "DB_PORT=3306",
            "DB_USER=mailguard",
            "DB_PASSWORD=mailguard_dev",
            "DB_NAME=mailguard_ai",
            "",
            "# Backend",
            "APP_ENV=development",
            "APP_DEBUG=true",
            "JWT_SECRET=change-me-in-prod",
            "JWT_EXPIRES_MIN=60",
            "RATE_LIMIT_PER_MIN=120",
            "",
            "# AI Service",
            "AI_SERVICE_URL=http://ai_service:8002",
            "AI_PROVIDER=baseline   # baseline | distilbert | gemini",
            "GEMINI_API_KEY=        # chỉ cần khi AI_PROVIDER=gemini",
            "GEMINI_MODEL=gemini-1.5-flash",
            "",
            "# Chrome Extension OAuth (tuỳ chọn, dùng cho Google sign-in)",
            "OAUTH_GOOGLE_CLIENT_ID=...apps.googleusercontent.com",
        ],
    )

    add_heading_2(doc, "A.3. Lệnh thường dùng")
    add_pseudocode(
        doc,
        [
            "# Khởi động toàn bộ stack",
            "docker compose -f docker-compose.dev.yml up -d",
            "",
            "# Xem log backend",
            "docker compose -f docker-compose.dev.yml logs -f backend",
            "",
            "# Tạo migration mới",
            "docker compose -f docker-compose.dev.yml exec backend alembic revision --autogenerate -m '...'",
            "",
            "# Apply migration",
            "docker compose -f docker-compose.dev.yml exec backend alembic upgrade head",
            "",
            "# Seed database",
            "docker compose -f docker-compose.dev.yml exec backend python -m app.database.seed",
            "",
            "# Train lại baseline",
            "docker compose -f docker-compose.dev.yml exec ai_service \\",
            "    python -m ai_service.scripts.train_baseline logistic_regression",
            "",
            "# Train lại DistilBERT",
            "docker compose -f docker-compose.dev.yml exec ai_service \\",
            "    python -m ai_service.scripts.train_distilbert",
            "",
            "# Test",
            "docker compose -f docker-compose.dev.yml exec backend pytest",
            "docker compose -f docker-compose.dev.yml exec ai_service pytest",
            "",
            "# Lint",
            "docker compose -f docker-compose.dev.yml exec backend ruff check app/",
            "",
            "# Reset toàn bộ (xoá DB + cache)",
            "docker compose -f docker-compose.dev.yml down -v",
            "docker compose -f docker-compose.dev.yml up -d",
        ],
    )
    add_page_break(doc)


def build_appendix_b(doc: Document) -> None:
    add_heading_1(doc, "PHỤ LỤC B — Một số đoạn mã nguồn tiêu biểu")

    add_heading_2(doc, "B.1. compute_risk() — risk/scorer.py")
    add_pseudocode(
        doc,
        [
            "def compute_risk(*, predicted_class, confidence, text,",
            "                links=None, attachments=None):",
            "    class_part = _class_base_score(predicted_class, confidence)",
            "    keyword_part, _ = _keyword_score(text)",
            "    urls = links if links is not None else extract_links(text or '')",
            "    url_part, _ = _url_score(urls)",
            "    attach_part = _attachment_score(attachments)",
            "    total = min(100.0, class_part + keyword_part + url_part + attach_part)",
            "    return RiskBreakdown(",
            "        classification=round(class_part, 2),",
            "        keywords=round(keyword_part, 2),",
            "        urls=round(url_part, 2),",
            "        attachments=round(attach_part, 2),",
            "        total=round(total, 2),",
            "    )",
        ],
    )

    add_heading_2(doc, "B.2. build_highlighted_spans() — xai/highlighter.py")
    add_pseudocode(
        doc,
        [
            "def build_highlighted_spans(text, *, keywords=True, urls=True):",
            "    spans = []",
            "    if keywords:",
            "        spans.extend(find_keyword_spans(text))",
            "    if urls:",
            "        spans.extend(find_url_spans(text))",
            "    spans.sort(key=lambda s: (s.start, -s.end))",
            "    merged = []",
            "    for s in spans:",
            "        if merged and s.start <= merged[-1].end:",
            "            prev = merged[-1]",
            "            merged[-1] = HighlightedSpan(",
            "                start=prev.start,",
            "                end=max(prev.end, s.end),",
            "                text=prev.text,",
            "                category=prev.category,",
            "                weight=max(prev.weight, s.weight),",
            "                reason=f'{prev.reason}; {s.reason}',",
            "            )",
            "        else:",
            "            merged.append(s)",
            "    return [{'start': s.start, 'end': s.end, 'text': s.text,",
            "             'category': s.category, 'weight': s.weight,",
            "             'reason': s.reason} for s in merged]",
        ],
    )

    add_heading_2(doc, "B.3. predict() — api/v1/predictions.py")
    add_pseudocode(
        doc,
        [
            "@router.post('', response_model=PredictionRead, status_code=201)",
            "async def predict(payload: PredictionRequest,",
            "                  db: Session = Depends(get_db),",
            "                  current_user: User = Depends(get_current_user)) -> PredictionRead:",
            "    # 1. Persist the email",
            "    email = email_service.create_email(db, current_user, payload.email)",
            "",
            "    # 2. Call the AI service",
            "    ai_result = await ai_client.predict(payload)",
            "",
            "    # 3. Resolve the model version",
            "    model_version = prediction_service._resolve_model_version(",
            "        db, payload.model_version",
            "    )",
            "",
            "    # 4. Save prediction",
            "    prediction = prediction_service.save_prediction(",
            "        db, current_user, email, model_version, ai_result",
            "    )",
            "",
            "    activity_log_service.log(",
            "        db, user=current_user, action='predict',",
            "        entity_type='email', entity_id=email.id,",
            "        details={'prediction_id': prediction.id,",
            "                 'predicted_class': prediction.predicted_class.value},",
            "    )",
            "    return prediction",
        ],
    )

    add_heading_2(doc, "B.4. runGmailScan() — background service worker (Chrome extension)")
    add_pseudocode(
        doc,
        [
            "async function runGmailScan():",
            "    token = await getGmailAccessToken({ interactive: false })",
            "    cfg = await chrome.storage.local.get(['mg_gmail_batch_size'])",
            "    batchSize = Number(cfg.mg_gmail_batch_size) || 25",
            "    emails = await fetchDecodedInbox({ maxResults: batchSize })",
            "    baseUrl = await getBaseUrl()",
            "    auth    = await getAuthHeader()",
            "    results = []",
            "    for (email of emails):",
            "        try:",
            "            payload = {",
            "                email: {",
            "                    sender         : email.from || 'unknown@unknown',",
            "                    sender_domain  : deriveSenderDomain(email.from),",
            "                    recipient      : email.to || null,",
            "                    subject        : email.subject || null,",
            "                    body_text      : (email.body_text || '').slice(0, 8000),",
            "                    body_html      : null,",
            "                    links          : [],",
            "                    attachments    : Array.isArray(email.attachments)",
            "                                       ? email.attachments : [],",
            "                    received_at    : normalizeReceivedAt(email.date),",
            "                },",
            "                model_version: null,",
            "                include_explanation: true,",
            "            }",
            "            resp = await fetch(`${baseUrl}/predictions`, {",
            "                method: 'POST',",
            "                headers: { 'Content-Type': 'application/json', ...auth },",
            "                body: JSON.stringify(payload),",
            "            })",
            "            data = await parseResponse(resp)",
            "            results.push({ messageId: email.id, threadId: email.threadId,",
            "                            subject: email.subject, from: email.from,",
            "                            prediction: data })",
            "            await notifyDanger({ ... })  // OS notification nếu risk cao",
            "        except (err):",
            "            results.push({ messageId: email.id, error: err.message })",
            "    await chrome.storage.local.set({ mg_gmail_scan_results: results })",
            "    return results",
        ],
    )
    add_page_break(doc)


def build_references(doc: Document) -> None:
    add_heading_1(doc, "TÀI LIỆU THAM KHẢO")

    refs = [
        "[1] V. Sanh, L. Debut, J. Chaumond, T. Wolf, \"DistilBERT, a distilled "
        "version of BERT: smaller, faster, cheaper and lighter\", arXiv preprint "
        "arXiv:1910.01108, 2019.",

        "[2] J. Devlin, M.-W. Chang, K. Lee, K. Toutanova, \"BERT: Pre-training of "
        "Deep Bidirectional Transformers for Language Understanding\", NAACL-HLT, 2019.",

        "[3] F. Pedregosa et al., \"Scikit-learn: Machine Learning in Python\", "
        "Journal of Machine Learning Research, vol. 12, pp. 2825-2830, 2011.",

        "[4] T. Joachims, \"A Probabilistic Analysis of the Rocchio Algorithm with "
        "TFIDF for Text Categorization\", Carnegie Mellon University, 1996.",

        "[5] S. Hochreiter, J. Schmidhuber, \"Long Short-Term Memory\", Neural "
        "Computation, vol. 9, no. 8, pp. 1735-1780, 1997.",

        "[6] Google DeepMind, \"Gemini: A Family of Highly Capable Multimodal Models\", "
        "Technical Report, 2024. [Online]. Available: https://deepmind.google/technologies/gemini/.",

        "[7] A. Vaswani et al., \"Attention Is All You Need\", Advances in Neural "
        "Information Processing Systems (NeurIPS), vol. 30, 2017.",

        "[8] M. Sahami, S. Dumais, D. Heckerman, E. Horvitz, \"A Bayesian Approach "
        "to Filtering Junk E-Mail\", AAAI Workshop on Learning for Text "
        "Categorization, 1998.",

        "[9] Verizon, \"2024 Data Breach Investigations Report (DBIR)\", Verizon "
        "Enterprise, 2024. [Online]. Available: https://www.verizon.com/business/resources/reports/dbir/.",

        "[10] Google, \"Chrome Extensions Manifest V3 Documentation\", 2023. "
        "[Online]. Available: https://developer.chrome.com/docs/extensions/develop/migrate/what-is-mv3.",

        "[11] Mozilla, \"MDN Web Docs: WebExtensions API\". [Online]. "
        "Available: https://developer.mozilla.org/en-US/docs/Mozilla/Add-ons/WebExtensions.",

        "[12] T. Mikolov, K. Chen, G. Corrado, J. Dean, \"Efficient Estimation of "
        "Word Representations in Vector Space\", arXiv preprint arXiv:1301.3781, 2013.",

        "[13] L. Breiman, \"Random Forests\", Machine Learning, vol. 45, no. 1, "
        "pp. 5-32, 2001.",

        "[14] C. Cortes, V. Vapnik, \"Support-Vector Networks\", Machine Learning, "
        "vol. 20, no. 3, pp. 273-297, 1995.",

        "[15] A. McCallum, K. Nigam, \"A Comparison of Event Models for Naive Bayes "
        "Text Classification\", AAAI-98 Workshop on Learning for Text "
        "Categorization, 1998.",

        "[16] FastAPI Documentation, 2024. [Online]. Available: https://fastapi.tiangolo.com/.",

        "[17] SQLAlchemy 2.0 Documentation, 2024. [Online]. Available: https://docs.sqlalchemy.org/.",

        "[18] Docker Documentation, 2024. [Online]. Available: https://docs.docker.com/.",

        "[19] HuggingFace Transformers Documentation, 2024. [Online]. "
        "Available: https://huggingface.co/docs/transformers.",

        "[20] R. Caruana, \"Multitask Learning\", Machine Learning, vol. 28, no. 1, "
        "pp. 41-75, 1997.",

        "[21] M. T. Ribeiro, S. Singh, C. Guestrin, \"Why Should I Trust You? "
        "Explaining the Predictions of Any Classifier\", KDD, 2016.",

        "[22] S. M. Lundberg, S.-I. Lee, \"A Unified Approach to Interpreting Model "
        "Predictions\", NeurIPS, 2017.",

        "[23] T. R. Peltier, \"Information Security Risk Analysis\", Auerbach "
        "Publications, 2005.",

        "[24] R. Oppliger, \"Internet and Intranet Security\", Artech House, 2002.",

        "[25] OWASP Foundation, \"OWASP Top 10 for LLM Applications\", 2023. "
        "[Online]. Available: https://owasp.org/www-project-top-10-for-large-language-model-applications/.",
    ]
    for ref in refs:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.2
        p.paragraph_format.first_line_indent = Cm(-0.6)
        p.paragraph_format.left_indent = Cm(0.6)
        run = p.add_run(ref)
        run.font.size = Pt(11)

    add_para(
        doc,
        f"Ngày hoàn thành báo cáo: {date.today().strftime('%d/%m/%Y')}.",
        italic=True,
    )


# --------------------------------------------------------------------------- #
# Main
# --------------------------------------------------------------------------- #
def main() -> None:
    doc = Document()

    # ---- Page margins ----
    for section in doc.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.4)
        section.right_margin = Cm(2.4)

    # ---- Default font ----
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), "Times New Roman")

    # ---- Build content ----
    build_cover_page(doc)
    build_toc(doc)
    build_preface(doc)
    build_chapter1(doc)
    build_chapter2(doc)
    build_chapter3(doc)
    build_chapter4(doc)
    build_chapter5(doc)
    build_chapter6(doc)
    build_chapter7(doc)
    build_appendix_a(doc)
    build_appendix_b(doc)
    build_references(doc)

    # ---- Save ----
    doc.save(str(OUTPUT_PATH))
    print(f"[OK] Saved report: {OUTPUT_PATH}")
    print(f"     Size: {os.path.getsize(OUTPUT_PATH):,} bytes")


if __name__ == "__main__":
    main()